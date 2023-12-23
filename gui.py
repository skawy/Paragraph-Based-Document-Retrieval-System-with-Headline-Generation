import PySimpleGUI as sg

import pandas as pd

from docx import Document
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer


def read_pdf(file_path) :
    pdf_text = []

    for page in extract_pages(file_path):
        page_elements = [(element.y1, element) for element in page._objs]
        page_elements.sort(key=lambda a: a[0], reverse=True)
        for _ , element in page_elements:
            # Check if the element is text element
            if isinstance(element, LTTextContainer):
                # Use the function to extract the text for each text element
                line_text = element.get_text()
                pdf_text.append(line_text)
    return pdf_text

def read_document(file_path):
    """
    Iam Assuming That csv or xlsx Has Column called Text contains paragraph
    Row = 1 paragraph
    """

    if file_path.endswith('.csv'):
        return  pd.read_csv(file_path)
    
    elif file_path.endswith('.xlsx'):
        return  pd.read_excel(file_path)

    elif file_path.endswith('.txt'):
        with open(file_path) as file:
            file_text = file.read()
        paragraphs = file_text.split('\n\n')
        return pd.DataFrame(paragraphs, columns=['Text'])
    
    elif file_path.endswith('.doc'):
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        paragraphs_as_strings = [p.text for p in paragraphs]
        return  pd.DataFrame(paragraphs_as_strings, columns=['Text'])

    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        paragraphs_as_strings = [p.text for p in paragraphs]
        return  pd.DataFrame(paragraphs_as_strings, columns=['Text'])

    elif file_path.endswith('.pdf'):
        pdf_text = read_pdf(file_path)
        return  pd.DataFrame(pdf_text, columns=['Text'])
 
    else :
        return None



def gui():
    document_uploaded = [
        [sg.Text(size=(50, 3), key='-DOCUMENT-')],
        [sg.Text(size=(50, 3), key='-IsLoaded-')]
    ]

    document_browsing = [
        [sg.Text("Choose Document to Upload: "), sg.Input( enable_events= True , key = "-UPLOAD-"), sg.FileBrowse()],
        
    ]

    layout = [
        [
            sg.Column(document_uploaded),
            sg.VSeparator(),
            sg.Column(document_browsing),

        ]
    ]

    document_uploader=sg.Window('Document Upload', layout, finalize=True)

    while True:

        event, values = document_uploader.read()
        if event == sg.WIN_CLOSED or event=="Exit":
            break
        
        document_path = values["-UPLOAD-"]
        document_uploader['-DOCUMENT-'].update(document_path)
        if event == '-UPLOAD-':
            document_path = values["-UPLOAD-"]
            df = read_document(document_path) 
            if df is not None:
                document_uploader['-IsLoaded-'].update("Iam Extracted The DataFrame Successfully")
                print(df)


    document_uploader.close()

if __name__ == '__main__':
    gui()

